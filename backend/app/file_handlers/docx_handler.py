"""
DOCX 文件文本提取
使用 python-docx
"""
from pathlib import Path
from docx import Document


def extract_text_from_docx(file_path: str) -> str:
    """
    从 Word (.docx) 文件中提取纯文本
    :param file_path: DOCX 文件路径
    :return: 提取的文本内容
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX 文件不存在: {file_path}")

    text_parts = []
    try:
        doc = Document(str(path))

        # 提取段落文本
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # 提取表格中的文本（简历常用表格布局）
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

    except Exception as e:
        raise RuntimeError(f"DOCX 解析失败: {str(e)}")

    full_text = "\n".join(text_parts)
    if not full_text.strip():
        raise ValueError("DOCX 文件中未提取到文本内容")

    return full_text
